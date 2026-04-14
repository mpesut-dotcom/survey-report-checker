"""
Report generation — Word document output.
Aggregates verification results into structured reports.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from checker.models.verification import SlideVerification
from checker.utils.json_utils import sanitize_text


# ──────────────────────────────────────────────────────────────────────
# Color constants
# ──────────────────────────────────────────────────────────────────────
_RED = RGBColor(0xCC, 0x00, 0x00)
_ORANGE = RGBColor(0xFF, 0x88, 0x00)
_GREEN = RGBColor(0x00, 0x88, 0x00)
_GRAY = RGBColor(0x66, 0x66, 0x66)
_BLUE = RGBColor(0x00, 0x55, 0xAA)


def generate_report(
    verifications: list[SlideVerification],
    output_path: Path,
    *,
    pptx_name: str = "",
    cost_info: dict | None = None,
) -> Path:
    """Generate Word report from verification results."""
    doc = Document()

    # Title
    title = doc.add_heading("QC Report — Provjera izvještaja", level=0)
    if pptx_name:
        p = doc.add_paragraph()
        _add_run(p, f"Prezentacija: {pptx_name}", color=_GRAY, size=11)

    # Summary
    _add_summary_section(doc, verifications)

    # Cost info
    if cost_info:
        _add_cost_section(doc, cost_info)

    # Detailed findings
    doc.add_heading("Detaljni nalaz po slajdovima", level=1)

    for v in sorted(verifications, key=lambda x: x.slide_number):
        _add_slide_section(doc, v)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _add_summary_section(
    doc: Document,
    verifications: list[SlideVerification],
):
    """Add summary statistics section."""
    doc.add_heading("Sažetak", level=1)

    total = len(verifications)
    errors = sum(1 for v in verifications if v.overall_status == "error")
    warnings = sum(1 for v in verifications if v.overall_status == "warning")
    ok = sum(1 for v in verifications if v.overall_status == "ok")
    info = sum(1 for v in verifications if v.overall_status == "info")

    p = doc.add_paragraph()
    _add_run(p, f"Verificirano slajdova: {total}\n", size=11)
    _add_run(p, f"  ✓ OK: {ok}", color=_GREEN, size=11)
    _add_run(p, f"  ⚠ Upozorenja: {warnings}", color=_ORANGE, size=11)
    _add_run(p, f"  ✗ Greške: {errors}\n", color=_RED, size=11)
    if info:
        _add_run(p, f"  ℹ Info (bez podataka): {info}", color=_GRAY, size=11)

    # Count issues by type
    data_errors = sum(len([i for i in v.data_issues if i.severity == "error"])
                      for v in verifications)
    data_warnings = sum(len([i for i in v.data_issues if i.severity == "warning"])
                        for v in verifications)
    data_info = sum(len([i for i in v.data_issues if i.severity == "info"])
                    for v in verifications)
    text_issues = sum(len(v.text_issues) for v in verifications)
    visual_issues = sum(len(v.visual_issues) for v in verifications)

    pass1_candidates = sum(v.pass1_total_candidates for v in verifications)
    pass1_datasets = sum(v.pass1_total_datasets for v in verifications)
    pass1_confident = sum(v.pass1_confident_datasets for v in verifications)
    match_sources = sum(len(v.match_sources) for v in verifications)
    match_failures = sum(len(v.match_failures) for v in verifications)
    low_conf = sum(
        1
        for v in verifications
        for mf in v.match_failures
        if mf.reason == "low_confidence"
    )
    unresolved = sum(
        1
        for v in verifications
        for mf in v.match_failures
        if mf.reason == "unresolved_after_pass1"
    )
    fallback_resolved = sum(
        1
        for v in verifications
        for ms in v.match_sources
        if ms.included_via == "pass1" and not ms.resolved_by.startswith("index:")
    )

    p2 = doc.add_paragraph()
    _add_run(p2, "Ukupno nalaza:\n", bold=True, size=11)
    _add_run(p2, f"  Data errors: {data_errors}, warnings: {data_warnings}, info: {data_info}\n", size=10)
    _add_run(p2, f"  Tekstualni problemi: {text_issues}\n", size=10)
    _add_run(p2, f"  Vizualni problemi: {visual_issues}", size=10)

    p3 = doc.add_paragraph()
    _add_run(p3, "Match dijagnostika:\n", bold=True, size=11)
    _add_run(p3, f"  Pass1 kandidati ukupno: {pass1_candidates}\n", size=10)
    _add_run(p3, f"  Pass1 dataseti: {pass1_datasets}, confident: {pass1_confident}\n", size=10)
    _add_run(p3, f"  Korišteni izvori za Pass2: {match_sources}\n", size=10)
    _add_run(p3, f"  Match failovi: {match_failures} (low_conf: {low_conf}, unresolved: {unresolved})\n", size=10)
    _add_run(p3, f"  Resolve fallback (nije index): {fallback_resolved}", size=10)


def _add_cost_section(doc: Document, cost_info: dict):
    """Add API cost section."""
    p = doc.add_paragraph()
    _add_run(p, "API troškovi: ", bold=True, size=10)
    _add_run(p, f"~${cost_info.get('total_cost_usd', 0):.3f} "
                f"({cost_info.get('total_calls', 0)} poziva, "
                f"{cost_info.get('total_input_tokens', 0):,} in / "
                f"{cost_info.get('total_output_tokens', 0):,} out tokena)",
             color=_GRAY, size=10)


def _add_slide_section(doc: Document, v: SlideVerification):
    """Add section for one slide's verification results."""
    status_icon = {"ok": "✓", "warning": "⚠", "error": "✗"}.get(v.overall_status, "?")
    status_color = {"ok": _GREEN, "warning": _ORANGE, "error": _RED}.get(v.overall_status, _GRAY)

    heading = doc.add_heading(level=2)
    run = heading.add_run(f"Slajd {v.slide_number} — {status_icon} {v.overall_status.upper()}")
    run.font.color.rgb = status_color

    if v.summary:
        p = doc.add_paragraph()
        _add_run(p, sanitize_text(v.summary), italic=True, color=_GRAY, size=10)

    if v.match_sources or v.match_failures:
        doc.add_heading("Izvori podataka", level=3)

        if v.match_sources:
            for src in v.match_sources:
                p = doc.add_paragraph(style="List Bullet")
                conf_str = f", conf={src.confidence:.2f}" if src.confidence is not None else ""
                qcode = src.question_code or "?"
                _add_run(
                    p,
                    f"[SOURCE] {src.excel_id} ({qcode}) via {src.included_via}/{src.resolved_by}{conf_str}",
                    color=_BLUE,
                    size=9,
                )
                if src.question_text:
                    _add_run(p, f"\n   {sanitize_text(src.question_text)}", color=_GRAY, size=9)

        if v.match_failures:
            for mf in v.match_failures:
                p = doc.add_paragraph(style="List Bullet")
                conf_str = f", conf={mf.confidence:.2f}" if mf.confidence is not None else ""
                qcode = mf.question_code or "?"
                _add_run(
                    p,
                    f"[FAIL] id={mf.excel_id or '(prazno)'} q={qcode} reason={mf.reason}{conf_str}",
                    color=_ORANGE,
                    size=9,
                )
                if mf.description:
                    _add_run(p, f"\n   {sanitize_text(mf.description)}", color=_GRAY, size=9)

    # Data issues
    if v.data_issues:
        doc.add_heading("Podaci", level=3)
        for issue in v.data_issues:
            p = doc.add_paragraph(style="List Bullet")
            sev_color = {"error": _RED, "warning": _ORANGE}.get(issue.severity, _GRAY)
            _add_run(p, f"[{issue.severity.upper()}] ", color=sev_color, bold=True, size=10)
            _add_run(p, sanitize_text(issue.detail), size=10)
            if issue.slide_value is not None or issue.excel_value is not None:
                _add_run(p, f"\n   Slajd: {issue.slide_value} → Excel: {issue.excel_value}",
                         color=_GRAY, size=9)

    # Text issues
    if v.text_issues:
        doc.add_heading("Tekst", level=3)
        for issue in v.text_issues:
            p = doc.add_paragraph(style="List Bullet")
            sev_color = {"error": _RED, "warning": _ORANGE}.get(issue.severity, _GRAY)
            _add_run(p, f"[{issue.severity.upper()}] ", color=sev_color, bold=True, size=10)
            _add_run(p, sanitize_text(issue.detail), size=10)

    # Visual issues
    if v.visual_issues:
        doc.add_heading("Vizualno", level=3)
        for issue in v.visual_issues:
            p = doc.add_paragraph(style="List Bullet")
            sev_color = {"error": _RED, "warning": _ORANGE}.get(issue.severity, _GRAY)
            _add_run(p, f"[{issue.severity.upper()}] ", color=sev_color, bold=True, size=10)
            _add_run(p, sanitize_text(issue.detail), size=10)


def _add_run(paragraph, text: str, *, bold=False, italic=False,
             color=None, size=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(sanitize_text(text))
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run
